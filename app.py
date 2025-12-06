import streamlit as st
import streamlit.components.v1 as components
from playwright.sync_api import sync_playwright
from bs4 import BeautifulSoup, NavigableString, Tag, Comment
import time
import subprocess
import os
import re
from io import BytesIO

# Word作成用
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 設定
# ==========================================
FIXED_ENTRY_URL = "https://www.h-ken.net/mypage/20250611_1605697556/"

# ==========================================
# サーバー設定
# ==========================================
def install_playwright():
    try:
        subprocess.run(["playwright", "install", "chromium"], check=True)
    except Exception as e:
        print(f"Install error: {e}")

if "setup_done" not in st.session_state:
    with st.spinner("サーバー起動中..."):
        install_playwright()
        st.session_state.setup_done = True

# ==========================================
# 便利関数：ファイル名用クリーニング
# ==========================================
def sanitize_filename(text):
    """ファイル名に使えない文字を削除"""
    if not text:
        return "story"
    # 禁止文字を全角などに置換するか削除
    text = re.sub(r'[\\/*?:"<>|]', "", text)
    text = text.replace('\n', '').replace('\r', '').replace('\t', '')
    text = text.strip()
    if len(text) > 60: # 長すぎるとエラーになるのでカット
        text = text[:60]
    return text if text else "story"

# ==========================================
# 色解析ロジック（ブラウザ計算値利用）
# ==========================================
def get_rgb_from_str(color_str):
    if not color_str: return None
    c = color_str.lower().strip()
    
    # rgba(...)
    rgb_match = re.search(r'rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)', c)
    if rgb_match:
        return RGBColor(int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3)))

    # Hex
    hex_match = re.search(r'#([0-9a-f]{6})', c)
    if hex_match:
        h = hex_match.group(1)
        return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:], 16))
    
    # 基本マップ
    colors = {
        'red': RGBColor(255, 0, 0), 'blue': RGBColor(0, 0, 255), 'green': RGBColor(0, 128, 0),
        'black': RGBColor(0, 0, 0), 'white': RGBColor(255, 255, 255),
        'orange': RGBColor(255, 165, 0), 'pink': RGBColor(255, 192, 203),
        'purple': RGBColor(128, 0, 128), 'gold': RGBColor(255, 215, 0)
    }
    return colors.get(c.split()[0])

def apply_style_to_run(run, element):
    """データ属性(data-calc-color)を見てスタイル適用"""
    calc_color = element.get('data-calc-color')
    calc_bold = element.get('data-calc-bold')
    
    if calc_bold == 'true':
        run.bold = True
    elif element.name in ['b', 'strong', 'h1', 'h2']:
        run.bold = True
        
    if calc_color:
        rgb = get_rgb_from_str(calc_color)
        if rgb:
            run.font.color.rgb = rgb
            return

    style_attr = element.get('style', '').lower()
    if 'color' in style_attr:
        m = re.search(r'color\s*:\s*([^;"]+)', style_attr)
        if m: 
            rgb = get_rgb_from_str(m.group(1))
            if rgb: run.font.color.rgb = rgb

# ==========================================
# Word作成エンジン（空白行対応版）
# ==========================================
BLOCK_TAGS = ['p', 'div', 'h1', 'h2', 'h3', 'blockquote', 'li', 'article', 'section']

def process_node_recursive(paragraph, node):
    if isinstance(node, NavigableString):
        text = str(node)
        # 本文以外のシステムコメントを除外
        if "contents_within" not in text:
            # 空白だけのテキストも、改行の意味を持つことがあるので完全無視はしない
            # ただしWordでは連続する空白は無視されるため、意味のある文字があるか確認
            if text.strip():
                run = paragraph.add_run(text)
                if node.parent:
                    apply_style_to_run(run, node.parent)
                
    elif isinstance(node, Tag):
        if node.name == 'br':
            # <br> は確実に改行させる
            paragraph.add_run('\n')
        elif node.name in ['script', 'style', 'noscript']:
            pass
        else:
            # 子要素を処理
            for child in node.children:
                process_node_recursive(paragraph, child)
            
            # ブロック要素が終わったら改行を入れる
            # これにより <p>あ</p><p>い</p> がくっつかずに改行される
            if node.name in BLOCK_TAGS:
                paragraph.add_run('\n')

def create_rich_docx(title_html, body_html):
    doc = Document()
    
    # タイトル
    soup_title = BeautifulSoup(title_html, 'html.parser')
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    process_node_recursive(p_title, soup_title)
    
    for run in p_title.runs:
        run.font.size = Pt(16)
        if not run.bold: run.bold = True

    doc.add_paragraph("") 

    # 本文
    soup_body = BeautifulSoup(body_html, 'html.parser')
    
    # ルート直下の要素ごとに段落を作成する方式に変更
    # これにより、大きなブロック間の余白が自然になる
    top_level_elements = soup_body.find_all(True, recursive=False)
    
    if not top_level_elements:
        # 要素がない（テキスト直書きなど）場合は1つの段落で
        p = doc.add_paragraph()
        process_node_recursive(p, soup_body)
    else:
        for element in top_level_elements:
            p = doc.add_paragraph()
            # 行間を少し詰めたい場合はここを調整（デフォルトは広め）
            # p.paragraph_format.space_after = Pt(0) 
            
            process_node_recursive(p, element)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# ブラウザ操作
# ==========================================
def fetch_html_force_clean(target_url):
    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=True,
            args=['--no-sandbox', '--disable-dev-shm-usage', '--disable-gpu']
        )
        iphone_12 = p.devices['iPhone 12']
        context = browser.new_context(**iphone_12)
        page = context.new_page()

        try:
            page.goto(FIXED_ENTRY_URL, timeout=30000)
            time.sleep(2) 
            page.goto(target_url, timeout=30000)
            page.wait_for_load_state("domcontentloaded")
            time.sleep(2) 

            page.evaluate("""
                () => {
                    // ポップアップ破壊
                    const keywords = ['はい', 'YES', 'Yes', '18歳', 'Enter', '入り口', '入場'];
                    const buttons = document.querySelectorAll('a, button, div, span');
                    for (let btn of buttons) {
                        if (keywords.some(k => btn.innerText && btn.innerText.includes(k))) {
                            btn.click();
                        }
                    }
                    const allDivs = document.querySelectorAll('body > div, body > section');
                    allDivs.forEach(div => {
                        const style = window.getComputedStyle(div);
                        if (style.position === 'fixed' && style.zIndex > 50) {
                            div.remove();
                        }
                    });
                    document.body.style.overflow = 'visible';
                    document.body.style.height = 'auto';
                    
                    // 色情報焼き付け
                    const targetArea = document.getElementById('sentenceBox') || document.body;
                    const allElements = targetArea.querySelectorAll('*');
                    allElements.forEach(el => {
                        const style = window.getComputedStyle(el);
                        const color = style.color;
                        const weight = style.fontWeight;
                        if (color && color !== 'rgb(0, 0, 0)') {
                            el.setAttribute('data-calc-color', color);
                        }
                        if (weight === 'bold' || parseInt(weight) >= 700) {
                            el.setAttribute('data-calc-bold', 'true');
                        }
                    });
                }
            """)
            time.sleep(1) 
            return page.content()
        except Exception as e:
            st.error(f"エラー: {e}")
            return None
        finally:
            browser.close()

# ==========================================
# 抽出ロジック
# ==========================================
def extract_target_content(html_content, target_url):
    soup = BeautifulSoup(html_content, 'html.parser')

    styles = []
    for link in soup.find_all('link', rel='stylesheet'):
        styles.append(str(link))
    for style in soup.find_all('style'):
        styles.append(str(style))
    style_html = "\n".join(styles)

    # タイトル取得
    title_html = ""
    title_text_clean = "無題"
    
    target_h1 = soup.find("h1", class_="pageTitle")
    if target_h1:
        title_html = str(target_h1)
        title_text_clean = target_h1.get_text(strip=True)
    else:
        target_h1 = soup.find("h1")
        if target_h1:
            title_html = str(target_h1)
            title_text_clean = target_h1.get_text(strip=True)
    
    if title_text_clean == "無題" and soup.title:
        title_text_clean = soup.title.get_text(strip=True)

    body_html = "<div>本文が見つかりませんでした</div>"
    target_div = soup.find(id="sentenceBox")
    if not target_div:
        target_div = soup.find(id="main_txt")

    if target_div:
        for comment in target_div.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()
        for bad in target_div.find_all(["script", "noscript", "iframe", "form", "button", "input"]):
            bad.decompose()
        cut_point = target_div.find(class_="kakomiPop2")
        if cut_point:
            for sibling in cut_point.find_next_siblings():
                sibling.decompose()
            cut_point.decompose()
        bad_words = ["無断転載", "Googleに通報", "刑事告訴", "民事訴訟", "エチケン", "contents_within"]
        for tag in target_div.find_all(['p', 'div', 'span', 'font', 'b']):
            text = tag.get_text()
            if any(w in text for w in bad_words):
                if len(text) < 400:
                    tag.decompose()
        body_html = str(target_div)

    final_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <base href="{target_url}">
        {style_html}
        <style>
            body {{ background-color: #fff; padding: 15px; font-family: sans-serif; overflow: auto !important; }}
            h1.pageTitle {{ font-size: 20px; margin-bottom: 20px; border-bottom: 1px solid #ccc; padding-bottom: 10px; line-height: 1.4; }}
            #sentenceBox {{ font-size: 16px; line-height: 1.8; color: #333; }}
        </style>
    </head>
    <body>
        {title_html}
        {body_html}
    </body>
    </html>
    """

    return title_html, body_html, final_html, title_text_clean

# ==========================================
# 画面構成
# ==========================================
st.set_page_config(page_title="H-Review Ultra", layout="centered")

st.title("💎 究極版コンテンツ抽出")
st.caption("全色対応・空白行維持・ファイル名自動化")

url = st.text_input("読みたい記事のURL", placeholder="https://...")

if st.button("抽出を開始する", type="primary", use_container_width=True):
    if not url:
        st.warning("URLを入力してください。")
    else:
        status = st.empty()
        status.info("⏳ 解析中...")
        
        html = fetch_html_force_clean(url)

        if html:
            status.info("📄 データ生成中...")
            
            title_html_str, body_html_str, final_html_preview, title_text_clean = extract_target_content(html, url)
            
            status.empty()
            st.success("完了！")
            
            # Word作成
            docx_file = create_rich_docx(title_html_str, body_html_str)
            
            # ファイル名設定
            safe_filename = sanitize_filename(title_text_clean) + ".docx"
            
            st.download_button(
                label=f"📘 「{safe_filename}」で保存",
                data=docx_file,
                file_name=safe_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.divider()
            components.html(final_html_preview, height=800, scrolling=True)
            
        else:
            status.error("読み込みに失敗しました。")
